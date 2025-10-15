# =============================================================================
# --- ARQUIVO: report_logic.py ---
# (Lógica para a geração de relatórios em Excel, PDF e Word)
# =============================================================================

import pandas as pd
from tkinter import filedialog, messagebox
import os
from datetime import datetime

# --- Tenta importar as bibliotecas de PDF e Word ---
try:
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def export_clients_to_excel(clients_list, parent_window=None):
    """
    Recebe uma lista de dicionários de clientes e a exporta para um arquivo Excel.
    """
    if not clients_list:
        messagebox.showwarning("Nenhum Dado", "Não há clientes na lista para exportar.", parent=parent_window)
        return

    try:
        data_for_export = {
            "Nome do Cliente": [client.get('name', 'N/A') for client in clients_list],
            "Código": [client.get('code', 'N/A') for client in clients_list]
        }
        df = pd.DataFrame(data_for_export)
    except Exception as e:
        messagebox.showerror("Erro ao Preparar Dados", f"Houve um problema ao organizar os dados:\n{e}", parent=parent_window)
        return

    timestamp = datetime.now().strftime("%Y-%m-%d")
    initial_filename = f"Relatorio_Codigos_de_Clientes_{timestamp}.xlsx"
    
    filepath = filedialog.asksaveasfilename(
        title="Salvar Relatório de Clientes em Excel",
        parent=parent_window, initialfile=initial_filename,
        defaultextension=".xlsx", filetypes=[("Arquivos Excel", "*.xlsx")]
    )

    if not filepath: return

    try:
        writer = pd.ExcelWriter(filepath, engine='openpyxl')
        df.to_excel(writer, index=False, sheet_name='Clientes')
        worksheet = writer.sheets['Clientes']
        for column_cells in worksheet.columns:
            max_length = max(len(str(cell.value)) for cell in column_cells)
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column_cells[0].column_letter].width = adjusted_width
        writer.close()

        if messagebox.askyesno("Sucesso", f"Relatório salvo com sucesso!\n\nDeseja abrir o arquivo agora?", parent=parent_window):
            os.startfile(filepath)
            
    except Exception as e:
        messagebox.showerror("Erro ao Salvar", f"Ocorreu um erro ao salvar o arquivo Excel:\n\n{e}", parent=parent_window)

def export_clients_to_pdf(clients_list, parent_window=None):
    """Gera um relatório em PDF com a lista de clientes."""
    if not REPORTLAB_AVAILABLE:
        messagebox.showerror("Biblioteca Faltando", "A biblioteca 'reportlab' é necessária para gerar PDFs.\nInstale com: pip install reportlab", parent=parent_window)
        return
        
    if not clients_list:
        messagebox.showwarning("Nenhum Dado", "Não há clientes na lista para exportar.", parent=parent_window)
        return
        
    timestamp = datetime.now().strftime("%Y-%m-%d")
    initial_filename = f"Relatorio_Codigos_de_Clientes_{timestamp}.pdf"
    filepath = filedialog.asksaveasfilename(
        title="Salvar Relatório de Clientes em PDF",
        parent=parent_window, initialfile=initial_filename,
        defaultextension=".pdf", filetypes=[("Arquivos PDF", "*.pdf")]
    )
    if not filepath: return

    try:
        doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=30, bottomMargin=30)
        elements = []
        styles = getSampleStyleSheet()
        
        title = Paragraph("Relatório de Códigos de Cliente", styles['h1'])
        elements.append(title)
        elements.append(Spacer(1, 20))

        table_data = [['Nome do Cliente', 'Código']]
        for client in clients_list:
            table_data.append([
                Paragraph(client.get('name', ''), styles['Normal']), 
                Paragraph(client.get('code', ''), styles['Normal'])
            ])
        
        table = Table(table_data, colWidths=[300, 150])
        style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4e5d6c")),
            ('TEXTCOLOR',(0,0),(-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#f0f0f0")),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('TOPPADDING', (0,1), (-1,-1), 6),
            ('BOTTOMPADDING', (0,1), (-1,-1), 6),
        ])
        table.setStyle(style)
        
        elements.append(table)
        doc.build(elements)
        
        if messagebox.askyesno("Sucesso", f"Relatório PDF salvo com sucesso!\n\nDeseja abrir o arquivo agora?", parent=parent_window):
            os.startfile(filepath)
            
    except Exception as e:
        messagebox.showerror("Erro ao Salvar PDF", f"Ocorreu um erro ao salvar o arquivo PDF:\n\n{e}", parent=parent_window)

def export_clients_to_word(clients_list, parent_window=None):
    """Gera um relatório em Word (.docx) com a lista de clientes."""
    if not DOCX_AVAILABLE:
        messagebox.showerror("Biblioteca Faltando", "A biblioteca 'python-docx' é necessária para gerar arquivos Word.\nInstale com: pip install python-docx", parent=parent_window)
        return

    if not clients_list:
        messagebox.showwarning("Nenhum Dado", "Não há clientes na lista para exportar.", parent=parent_window)
        return

    timestamp = datetime.now().strftime("%Y-%m-%d")
    initial_filename = f"Relatorio_Codigos_de_Clientes_{timestamp}.docx"
    filepath = filedialog.asksaveasfilename(
        title="Salvar Relatório de Clientes em Word",
        parent=parent_window, initialfile=initial_filename,
        defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")]
    )
    if not filepath: return
    
    try:
        document = docx.Document()
        document.add_heading('Relatório de Códigos de Cliente', level=1)
        document.add_paragraph(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nome do Cliente'
        hdr_cells[1].text = 'Código'
        
        for client in clients_list:
            row_cells = table.add_row().cells
            row_cells[0].text = client.get('name', '')
            row_cells[1].text = client.get('code', '')
            
        document.save(filepath)
        
        if messagebox.askyesno("Sucesso", f"Relatório Word salvo com sucesso!\n\nDeseja abrir o arquivo agora?", parent=parent_window):
            os.startfile(filepath)

    except Exception as e:
        messagebox.showerror("Erro ao Salvar Word", f"Ocorreu um erro ao salvar o arquivo Word:\n\n{e}", parent=parent_window)